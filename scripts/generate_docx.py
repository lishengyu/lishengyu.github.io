"""
Generate GitHub Pages Setup Guide as a Word document (.docx)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ---- Set page margins (1 inch = 1440 twips = 1 inch in DXA, Word uses EMU: 1 inch = 914400 EMU) ----
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ---- Styles ----
style_normal = doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(11)

# Heading 1
h1 = doc.styles.add_style('MyHeading1', WD_STYLE_TYPE.PARAGRAPH)
h1.base_style = doc.styles['Normal']
h1.font.name = 'Arial'
h1.font.size = Pt(18)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
h1.paragraph_format.space_before = Pt(16)
h1.paragraph_format.space_after  = Pt(8)

# Heading 2
h2 = doc.styles.add_style('MyHeading2', WD_STYLE_TYPE.PARAGRAPH)
h2.base_style = doc.styles['Normal']
h2.font.name = 'Arial'
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after  = Pt(6)

# Heading 3
h3 = doc.styles.add_style('MyHeading3', WD_STYLE_TYPE.PARAGRAPH)
h3.base_style = doc.styles['Normal']
h3.font.name = 'Arial'
h3.font.size = Pt(12)
h3.font.bold = True
h3.paragraph_format.space_before = Pt(8)
h3.paragraph_format.space_after  = Pt(4)

def add_code_block(paragraph_text_lines):
    """Add a styled code block."""
    for line in paragraph_text_lines:
        p = doc.add_paragraph(line)
        p.style = 'Normal'
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        p.paragraph_format.left_indent  = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

# ============================================================
# Cover / Title Page
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GitHub Pages 个人网站')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('搭建完整指南')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('技术博客 · 简洁现代风格 · 免费托管')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2024')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.add_page_break()

# ============================================================
# Table of Contents
# ============================================================
doc.add_heading('目录', level=1)
toc_items = [
    ('1. 什么是 GitHub Pages',    '1'),
    ('2. 前期准备',               '1'),
    ('3. 创建 GitHub Pages 仓库', '2'),
    ('4. 部署网站文件',           '3'),
    ('5. 自定义域名（可选）',     '4'),
    ('6. 更新与维护',             '5'),
    ('7. 常见问题',               '6'),
]
for title, page in toc_items:
    p = doc.add_paragraph(f'{title} .......................... {page}')
    p.style = 'Normal'
doc.add_page_break()

# ============================================================
# Section 1
# ============================================================
doc.add_heading('1. 什么是 GitHub Pages', level=1)
doc.add_paragraph(
    'GitHub Pages 是 GitHub 提供的免费静态网站托管服务，'
    '支持用户站点（username.github.io）和项目站点两种使用方式。'
)
doc.add_paragraph('主要优点：')
for adv in ['完全免费，无流量限制', '原生支持静态站点生成器', '支持自定义域名 + 免费 HTTPS', '通过 Git 管理，版本控制完善', '无需服务器维护']:
    doc.add_paragraph(adv, style='List Bullet')

# ============================================================
# Section 2
# ============================================================
doc.add_heading('2. 前期准备', level=1)
doc.add_heading('2.1 所需账号与工具', level=2)
tools = [
    ('GitHub 账号', 'https://github.com'),
    ('Git',         'https://git-scm.com'),
    ('VS Code',     'https://code.visualstudio.com'),
]
for name, url in tools:
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{name}：{url}')

doc.add_heading('2.2 验证 Git 安装', level=2)
doc.add_paragraph('打开终端，运行以下命令验证 Git 是否安装成功：')
add_code_block(['git --version', '# 输出类似：git version 2.40.0'])

# ============================================================
# Section 3
# ============================================================
doc.add_heading('3. 创建 GitHub Pages 仓库', level=1)
doc.add_paragraph('按照以下步骤在 GitHub 上创建 Pages 仓库：')
steps = [
    '登录 GitHub，点击右上角 ＋ → New repository',
    '仓库名称填写：yourusername.github.io（替换 yourusername 为你的 GitHub 用户名）',
    '设置为 Public',
    '勾选 Add a README file',
    '点击 Create repository',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}', style='List Number')

warn = doc.add_paragraph()
warn.add_run('⚠️ 注意：').bold = True
warn.add_run('仓库名必须严格为 username.github.io，否则无法生效！')

# ============================================================
# Section 4
# ============================================================
doc.add_heading('4. 部署网站文件', level=1)
doc.add_heading('4.1 克隆仓库到本地', level=2)
add_code_block([
    'git clone https://github.com/yourusername/yourusername.github.io.git',
    'cd yourusername.github.io',
])

doc.add_heading('4.2 复制网站文件', level=2)
doc.add_paragraph('将网站模板文件复制到仓库目录，目录结构如下：')
add_code_block([
    'yourusername.github.io/',
    '├── index.html          ← 主入口',
    '├── css/',
    '│   └── style.css',
    '├── js/',
    '│   ├── data.js         ← 内容数据（编辑此文件自定义内容）',
    '│   └── app.js',
    '├── assets/',
    '│   └── images/',
    '└── README.md',
])

doc.add_heading('4.3 自定义内容', level=2)
doc.add_paragraph('编辑 js/data.js 来更新网站内容。以下为常用配置说明：')

doc.add_paragraph('① 修改博客文章（blogPosts 数组）：', style='Intense Quote')
add_code_block([
    'const blogPosts = [',
    '    {',
    '        title: "你的文章标题",',
    '        excerpt: "文章摘要...",',
    '        date: "2024-03-15",',
    '        tags: ["前端", "React"],',
    '        url: "#"',
    '    },',
    '];',
])

doc.add_paragraph('② 修改项目展示（projects 数组）：', style='Intense Quote')
add_code_block([
    'const projects = [',
    '    {',
    '        title: "项目名称",',
    '        description: "项目描述...",',
    '        icon: "🚀",',
    '        tags: ["Vue.js", "Node.js"],',
    '        demoUrl: "https://...",',
    '        codeUrl: "https://github.com/..."',
    '    },',
    '];',
])

doc.add_heading('4.4 提交并推送', level=2)
add_code_block([
    'git add .',
    'git commit -m "Initial commit: add personal website"',
    'git push origin main',
])

doc.add_paragraph('推送成功后，等待 1~2 分钟，访问 https://yourusername.github.io 即可查看网站。')

# ============================================================
# Section 5
# ============================================================
doc.add_heading('5. 自定义域名（可选）', level=1)
doc.add_paragraph('如果你想使用自己的域名（如 yourname.com），按照以下步骤配置：')

doc.add_heading('5.1 配置 DNS 解析', level=2)
doc.add_paragraph('在域名注册商管理后台添加 A 记录，指向 GitHub Pages 的 IP 地址：')

# DNS table
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['类型', '名称', '值']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

dns_data = [
    ('A',     '@', '185.199.108.153'),
    ('A',     '@', '185.199.109.153'),
    ('A',     '@', '185.199.110.153'),
    ('A',     '@', '185.199.111.153'),
]
for i, row_data in enumerate(dns_data, 1):
    for j, val in enumerate(row_data):
        table.rows[i].cells[j].text = val

doc.add_heading('5.2 在 GitHub 中配置', level=2)
cf_steps = [
    '进入仓库 Settings → Pages',
    '在 Custom domain 输入你的域名，点击 Save',
    '勾选 Enforce HTTPS（等待 DNS 生效后可勾选）',
]
for step in cf_steps:
    doc.add_paragraph(step, style='List Number')

# ============================================================
# Section 6
# ============================================================
doc.add_heading('6. 更新与维护', level=1)
doc.add_paragraph('网站上线后，日常更新非常简便：')
maintenance = [
    ('添加新博客文章', '在 js/data.js 的 blogPosts 数组开头添加新文章对象，然后 git push'),
    ('添加新项目',     '在 js/data.js 的 projects 数组开头添加新项目对象，然后 git push'),
    ('修改颜色主题',   '编辑 css/style.css 中的 CSS 变量（--color-primary 等）'),
    ('修改个人信息',   '编辑 index.html 中的姓名、简介、社交链接等'),
]
table = doc.add_table(rows=len(maintenance)+1, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['操作', '说明']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
for i, (op, desc) in enumerate(maintenance, 1):
    table.rows[i].cells[0].text = op
    table.rows[i].cells[1].text = desc

# ============================================================
# Section 7
# ============================================================
doc.add_heading('7. 常见问题', level=1)
faqs = [
    ('网站访问 404 怎么办？',
     '确认仓库名为 username.github.io；确认分支为 main；确认 index.html 在仓库根目录。'),
    ('推送后网站没有更新？',
     'GitHub Pages 有 1~3 分钟缓存延迟；查看仓库 Actions 标签页确认构建状态。'),
    ('如何添加自己的头像图片？',
     '将头像图片放到 assets/images/avatar.jpg，然后在 index.html 中替换字母头像为 <img> 标签。'),
    ('能否添加评论功能？',
     '可以使用 Giscus（基于 GitHub Discussions，免费）或 Utterances 添加评论功能。'),
    ('如何添加 Google Analytics？',
     '在 index.html 的 </head> 前添加 GA 跟踪代码即可。'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f'Q: {q}')
    run.bold = True
    run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
    doc.add_paragraph(a)
    doc.add_paragraph()

# ============================================================
# Appendix: Git Cheat Sheet
# ============================================================
doc.add_page_break()
doc.add_heading('附录：常用 Git 命令速查', level=1)
git_cmds = [
    ('git status',           '查看文件状态'),
    ('git add .',            '添加所有文件到暂存区'),
    ('git add index.html',   '添加指定文件'),
    ('git commit -m "msg"',  '提交变更'),
    ('git push origin main', '推送到远程仓库'),
    ('git pull origin main',  '拉取最新代码'),
    ('git log --oneline',    '查看提交历史'),
]
table = doc.add_table(rows=len(git_cmds)+1, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['命令', '说明']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
for i, (cmd, desc) in enumerate(git_cmds, 1):
    cell_cmd = table.rows[i].cells[0]
    cell_cmd.text = cmd
    for p in cell_cmd.paragraphs:
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    table.rows[i].cells[1].text = desc

# ---- Save ----
output_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'GitHub Pages 搭建指南.docx')
output_path = os.path.abspath(output_path)
doc.save(output_path)
print(f'✅ Word 文档已生成：{output_path}')
